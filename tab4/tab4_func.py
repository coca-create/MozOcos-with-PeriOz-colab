import gradio as gr
import os
from docx import Document
import codecs

def read_file_content(file):
    if file is None:
        return """<div style='color: orange !important; font-family: inherit; text-align: center; 
                display: flex; align-items: flex-start; justify-content: center; height: 400px; padding-top: 40px;'>
                No file uploaded
                </div>"""

    file_extension = os.path.splitext(file.name)[1]
    content = ""
    
    if file_extension == '.docx':
        doc = Document(file.name)
        content = "\n".join([para.text for para in doc.paragraphs])
        content = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 400px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{content}</pre>"""

    elif file_extension == '.txt':
        with codecs.open(file.name, 'r', 'utf-8') as f:
            content = f.read()
            content = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 500px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{content}</pre>"""

    elif file_extension == '.srt':
        with codecs.open(file.name, 'r', 'utf-8') as f:
            content = f.read()
            content = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 500px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{content}</pre>"""
    
    return content

def display_file_content(file):
    if file is None:
        return read_file_content(file), gr.update(value=""), gr.update(value=None)
    return read_file_content(file), gr.update(), gr.update()

def save_translated_content(file, translated_text):
    if file==None:
        return []
    file_name, file_extension = os.path.splitext(file.name)
    output_file_path = file_name + "_ja" + file_extension

    if file_extension == '.docx':
        doc = Document()
        doc.add_paragraph(translated_text)
        doc.save(output_file_path)

    elif file_extension == '.txt' or file_extension == '.srt':
        with codecs.open(output_file_path, 'w', 'utf-8') as f:
            f.write(translated_text)

    return output_file_path

def translate(file, translated_text):
    return save_translated_content(file, translated_text)
