import os
import tempfile
import json
from pydub import AudioSegment
import gradio as gr
from faster_whisper import WhisperModel
import zipfile
import torch
from docx import Document
import re
from openpyxl.styles import Alignment, Font, PatternFill
import pandas as pd

def get_audio_duration(File):
    try:
        file_path_str = File.name if hasattr(File, 'name') else File
        if not os.path.exists(file_path_str):
            return "File does not exist"
        audio = AudioSegment.from_file(file_path_str)
        duration = len(audio) / 1000.0
        return duration
    except AttributeError as e:
        return "AttributeError occurred"
    except Exception as e:
        return "An error occurred"

def format_timestamp(seconds):
    hrs, secs = divmod(seconds, 3600)
    mins, secs = divmod(secs, 60)
    millis = int((secs % 1) * 1000)
    return f"{int(hrs):02}:{int(mins):02}:{int(secs):02},{millis:03}"

def transcribe(File, Model, Computing, Lang, BeamSize, VadFilter, progress=gr.Progress()):
    
    if not File:
        error_message = "エラー: ファイルが提供されていません。"
        return error_message, "", "", [], [], "", "", "", "", ""

    try:
        FileName = File.name if hasattr(File, 'name') else File
        if Lang == "日本語":
            Lang = "ja"
        else:
            Lang = "en"
        
        model = WhisperModel(Model, device="cuda", compute_type=Computing)
        segments, _ = model.transcribe(File, word_timestamps=True, beam_size=BeamSize, initial_prompt="Hello, I am Scott.", language=Lang, vad_filter=VadFilter)
    except Exception as e:
        error_message = f"文字起こし中にエラーが発生しました: {e}"
        return error_message, "", "", [], [], "", "", "", "", ""

    total_duration = get_audio_duration(File)
    if isinstance(total_duration, str):  # get_audio_duration関数がエラーメッセージを返した場合
        return total_duration, "", "", [], [], "", "", "", "", ""
    words_data = []

    for segment in segments:
        for word in segment.words:
            word_info = {
                "start": word.start,
                "end": word.end,
                "word": word.word
            }
            words_data.append(word_info)
        progress(segment.end / total_duration)
        
     # メモリ解放
    del model
    torch.cuda.empty_cache()

    for word_info in words_data:
        word_info["word"] = word_info["word"].replace(" Dr.", " Dr★").replace(" dr.", " dr★")

    # 前処理: words_data内の各wordの中から★を削除する
    cleaned_words_data = []
    for word_info in words_data:
        cleaned_word_info = {
            "start": word_info["start"],
            "end": word_info["end"],
            "word": word_info["word"].replace("★", "")
        }
        cleaned_words_data.append(cleaned_word_info)

    input_file_name = os.path.splitext(os.path.basename(File))[0]
    temp_dir = tempfile.gettempdir()
    
    json_output_file_name = f"{input_file_name}.json"
    json_output_path = os.path.join(temp_dir, json_output_file_name)
    # JSONファイルへの書き込み
    with open(json_output_path, 'w', encoding='utf-8') as f:
       json.dump(cleaned_words_data, f, ensure_ascii=False, indent=4)
       
    # 書き込んだJSONデータの表示（デバッグ用）
    json_content = json.dumps(cleaned_words_data, ensure_ascii=False, indent=4)


    srt_entries = []
    entry_number = 1
    segment_text = ""
    segment_start = None
    segment_end = None

    for word_info in words_data:
        if segment_start is None:
            segment_start = word_info["start"]
        
        segment_text += word_info["word"]
        segment_end = word_info["end"]
        
        if word_info["word"].endswith('.'):
            srt_entries.append({
                "number": entry_number,
                "start": segment_start,
                "end": segment_end,
                "text": segment_text.strip()
            })
            entry_number += 1
            segment_text = ""
            segment_start = None

    if segment_text.strip():
        srt_entries.append({
            "number": entry_number,
            "start": segment_start,
            "end": segment_end,
            "text": segment_text.strip()
        })

    srt_output_file_name = f"{input_file_name}.srt"
    srt_output_path = os.path.join(temp_dir, srt_output_file_name)

    with open(srt_output_path, 'w', encoding='utf-8') as f:
        for entry in srt_entries:
            start_time = format_timestamp(entry["start"])
            end_time = format_timestamp(entry["end"])
            text = entry['text'].replace(" Dr★", " Dr.").replace(" dr★", " dr.").replace("Dr★", "Dr.")
            f.write(f"{entry['number']}\n{start_time} --> {end_time}\n{text}\n\n")

    with open(srt_output_path, 'r', encoding='utf-8') as f:
        srt_content = f.read()

    txt_nr_content = ""
    for word_info in words_data:
        if not txt_nr_content:
            txt_nr_content += word_info['word'].lstrip()
        else:
            txt_nr_content += word_info['word']

    txt_nr_output_file_name = f"{input_file_name}_NR.txt"
    txt_nr_output_path = os.path.join(temp_dir, txt_nr_output_file_name)
    with open(txt_nr_output_path, 'w', encoding='utf-8') as f:
        txt_nr_content = txt_nr_content.replace(" Dr★", " Dr.").replace(" dr★", " dr.").replace("Dr★", "Dr.")
        f.write(txt_nr_content)

    txt_r_content = ""
    previous_word_end = 0
    is_first_word = True
    for word in words_data:
        if is_first_word or txt_r_content.endswith("\n"):
            txt_r_content += word['word'].strip()
        else:
            txt_r_content += word['word']
    
        if "." in word['word']:
            if word['start'] - previous_word_end >= 0.5:
                txt_r_content += "\n"
            previous_word_end = word['end']
        is_first_word = False

    

    txt_r_output_file_name = f"{input_file_name}_R.txt"
    txt_r_output_path = os.path.join(temp_dir, txt_r_output_file_name)

    with open(txt_r_output_path, 'w', encoding='utf-8') as f:
        txt_r_content = txt_r_content.replace(" Dr★", " Dr.").replace(" dr★", " dr.").replace("Dr★", "Dr.")
        f.write(txt_r_content)

    # srtファイルからワードファイルへ変換
    doc_srt = Document()

    srtdoc_output_file_name = f"{input_file_name}_srt.docx"
    srtdoc_output_path = os.path.join(temp_dir, srtdoc_output_file_name)

    with open(srt_output_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    subtitle_number = None
    timestamp = None
    subtitle_text = []

    for line in lines:
        line = line.strip()
        if line.isdigit():
            # 以前の字幕エントリを追加
            if subtitle_number is not None and subtitle_text:
                doc_srt.add_paragraph(f'{subtitle_number}')
                doc_srt.add_paragraph(f'{timestamp}')
                doc_srt.add_paragraph(' '.join(subtitle_text))
                doc_srt.add_paragraph()  # 空行で区切る

            subtitle_number = line
            timestamp = None
            subtitle_text = []
        elif '-->' in line:
            timestamp = line
        elif line:
            subtitle_text.append(line)

    if subtitle_number is not None and subtitle_text:
        doc_srt.add_paragraph(f'{subtitle_number}')
        doc_srt.add_paragraph(f'{timestamp}')
        doc_srt.add_paragraph(' '.join(subtitle_text))

    doc_srt.save(srtdoc_output_path)

    ## txt(nr)をdoc変換
    txtdoc_nr = Document()
    txtdoc_nr_output_file_name = f"{input_file_name}_txtnr.docx"
    txtdoc_nr_output_path = os.path.join(temp_dir, txtdoc_nr_output_file_name)

    with open(txt_nr_output_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    for line in lines:
        txtdoc_nr.add_paragraph(line)

    txtdoc_nr.save(txtdoc_nr_output_path)

    ## txt(r)をdoc変換
    txtdoc_r = Document()
    txtdoc_r_output_file_name = f"{input_file_name}_txtr.docx"
    txtdoc_r_output_path = os.path.join(temp_dir, txtdoc_r_output_file_name)

    with open(txt_r_output_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    for line in lines:
        txtdoc_r.add_paragraph(line)

    txtdoc_r.save(txtdoc_r_output_path)


    # zipファイルにまとめる(srt,txtr,txtnr)。
    zip_core_file_name = f"{input_file_name}_core.zip"
    zip_core_file_path = os.path.join(temp_dir, zip_core_file_name)

    with zipfile.ZipFile(zip_core_file_path, 'w') as zip_file:
        zip_file.write(srt_output_path, os.path.basename(srt_output_path))
        zip_file.write(txt_r_output_path, os.path.basename(txt_r_output_path))
        zip_file.write(txt_nr_output_path, os.path.basename(txt_nr_output_path))


    # zipファイルにまとめる(doc)。
    zip_doc_file_name = f"{input_file_name}_docx_en.zip"
    zip_doc_file_path = os.path.join(temp_dir, zip_doc_file_name)

    with zipfile.ZipFile(zip_doc_file_path, 'w') as zip_file:
        zip_file.write(srtdoc_output_path,os.path.basename(srtdoc_output_path))
        zip_file.write(txtdoc_nr_output_path,os.path.basename(txtdoc_nr_output_path))
        zip_file.write(txtdoc_r_output_path,os.path.basename(txtdoc_r_output_path))

    print(f"Processed {FileName}")
    
    main_files = [srt_output_path,
                    txt_nr_output_path,
                    txt_r_output_path,
                    zip_core_file_path]
    
    zip_doc_file_path = os.path.join(temp_dir, zip_doc_file_name)

    doc_files=[srtdoc_output_path,txtdoc_nr_output_path,txtdoc_r_output_path,zip_doc_file_path]


    html_srt = f"""<pre style="white-space: pre-wrap; overflow-y: auto; height: 400px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{srt_content}</pre>"""
    html_nr_txt = f"""<pre style="white-space: pre-wrap; overflow-y:auto; height: 400px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{txt_nr_content}</pre>"""
    html_r_txt = f"""<pre style="white-space: pre-wrap; overflow-y:auto; height: 400px; word-wrap: break-word; padding: 10px; font-family: inherit; font-size: inherit;">{txt_r_content}</pre>"""

    '''global global_srt_file_path
    global_srt_file_path=srt_output_path'''
    filename_copy=input_file_name
    srt_dummy_output_path=srt_output_path
    return srt_content, txt_nr_content, txt_r_content, main_files, doc_files ,html_srt, html_nr_txt, html_r_txt,filename_copy,srt_dummy_output_path
