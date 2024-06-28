import gradio as gr
import docx
from docx import Document
import tempfile
import os
import re
import zipfile

def unify_timestamps_vtt(text):
    pattern_1_digit = re.compile(r'(\d{2}:\d{2}:\d{2}\.\d)(?!\d)')
    pattern_2_digits = re.compile(r'(\d{2}:\d{2}:\d{2}\.\d{2})(?!\d)')
    content = pattern_1_digit.sub(lambda x: x.group(1) + '00', text)
    content = pattern_2_digits.sub(lambda x: x.group(1) + '0', content)
    return content

def unify_timestamps(text):
    pattern_1_digit = re.compile(r'(\d{2}:\d{2}:\d{2},\d)(?!\d)')
    pattern_2_digits = re.compile(r'(\d{2}:\d{2}:\d{2},\d{2})(?!\d)')
    content = pattern_1_digit.sub(lambda x: x.group(1) + '00', text)
    content = pattern_2_digits.sub(lambda x: x.group(1) + '0', content)
    return content

def convert_docx_to_srttxt(docx_files):
    output_files = []
    if docx_files is None:
        return []
    for docx_file in docx_files:
        try:
            filename = os.path.basename(docx_file.name)
            base_name, ext = os.path.splitext(filename)
            clean_name = re.sub(r'[\r\n]+', '', base_name)
            print(f"Processing file: {clean_name}")

            if clean_name.endswith('_srt 1'):
                output_filename = clean_name.replace('_srt 1', '_ja.srt')
            elif clean_name.endswith("_srt"):
                output_filename = clean_name.replace("_srt", "_ja.srt")
            elif clean_name.endswith('_vtt 1'):
                output_filename = clean_name.replace('_vtt 1', '_ja.vtt')
            elif clean_name.endswith("_vtt"):
                output_filename = clean_name.replace("_vtt", "_ja.vtt")
            elif clean_name.endswith('_txtnr 1'):
                output_filename = clean_name.replace('_txtnr 1', '_NR_ja.txt')
            elif clean_name.endswith("_txtnr"):
                output_filename = clean_name.replace("_txtnr", "_NR_ja.txt")
            elif clean_name.endswith('_txtr 1'):
                output_filename = clean_name.replace('_txtr 1', '_R_ja.txt')
            elif clean_name.endswith("_txtr"):
                output_filename = clean_name.replace("_txtr", "_R_ja.txt")
            else:
                print(f"Skipping file with unrecognized pattern: {clean_name}")
                continue

            doc = Document(docx_file)
            content = "\n".join([para.text for para in doc.paragraphs])
            print(f"Initial content read from file: {content[:200]}")  # Show only the first 200 characters for brevity

            if clean_name.endswith("_srt") or clean_name.endswith("_srt 1"):
                content = unify_timestamps(content)
                content = re.sub(r'\s+', '', content)
                pattern = re.compile(r'(\d{1,4})\s*(\d{2}:\d{2}:\d{2},\d{3}\s*-->\s*\d{2}:\d{2}:\d{2},\d{3})')
            elif clean_name.endswith("_vtt") or clean_name.endswith("_vtt 1"):
                content = unify_timestamps_vtt(content)
                content = re.sub(r'\s+', '', content)
                pattern = re.compile(r'(\d{1,4})\s*(\d{2}:\d{2}:\d{2}\.\d{3}\s*-->\s*\d{2}:\d{2}:\d{2}\.\d{3})')

            print(f"Content after timestamp unification: {content[:200]}")  # Show only the first 200 characters for brevity

            segments = pattern.split(content)
            corrected_content = []

            for i in range(1, len(segments), 3):
                segment_id = segments[i]
                timestamp = segments[i + 1]
                text = segments[i + 2]
                
                corrected_content.append(f"{segment_id}")
                corrected_content.append(timestamp.replace('-->', ' --> '))
                corrected_content.append(text)

            final_content = "\n\n".join("\n".join(block) for block in zip(*[iter(corrected_content)] * 3))

            output_filepath = os.path.join(tempfile.gettempdir(), output_filename)
            with open(output_filepath, 'w', encoding='utf-8') as output_file:
                output_file.write(final_content)
            
            output_files.append(output_filepath)
        except Exception as e:
            print(f"An error occurred while processing {filename}: {str(e)}")
   

    if len(output_files)>1:
        zip_filename = os.path.join(tempfile.gettempdir(), "converted_from_docx_ja.zip")
        with zipfile.ZipFile(zip_filename, 'w') as zipf:
            for file in output_files:
                zipf.write(file, os.path.basename(file))
        
        output_files.append(zip_filename)
    
    return output_files



def clear_inputs():
    return None, None



def process_doc_files(files):
    output_files = []
    if files is None:
        return []
    for file in files:
        filename = os.path.basename(file.name)
        match = re.match(r"(.+?)(_NR\.txt|_R\.txt|\.srt|\.vtt)$", filename)
        if not match:
            continue  # skip files with unknown extensions
        
        basename, ext = match.groups()
        if ext == ".srt":
            doc_filename = f"{basename}_srt.docx"
        elif ext == ".vtt":
            doc_filename = f"{basename}_vtt.docx"
        elif ext == "_NR.txt":
            doc_filename = f"{basename}_txtnr.docx"
        elif ext == "_R.txt":
            doc_filename = f"{basename}_txtr.docx"

        if ext in ['.srt', '.vtt']:
            with open(file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if ext == '.srt':
                unified_content = unify_timestamps(content)
            else:
                unified_content = unify_timestamps_vtt(content)

            doc = Document()
            doc.add_paragraph(unified_content)
            doc.save(doc_filename)
            output_files.append(doc_filename)

        elif ext in ["_NR.txt", "_R.txt"]:
            with open(file.name, 'r', encoding='utf-8') as f:
                content = f.read()
            doc = Document()
            doc.add_paragraph(content)
            doc.save(doc_filename)
            output_files.append(doc_filename)
    
    if len(output_files) > 1:
        zip_filename = "converted_from_srttxt_en.zip"
        with zipfile.ZipFile(zip_filename, 'w') as zipf:
            for file in output_files:
                zipf.write(file)
        output_files.append(zip_filename)
    
    return output_files

def clear_both():
    return None, None


